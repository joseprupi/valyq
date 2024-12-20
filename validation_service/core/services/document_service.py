import os
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class DocumentService:
    """Service for generating documentation from test results"""
    
    def __init__(self, base_path: str):
        self.base_path = Path(base_path)
        
    def _parse_markdown_table(self, lines: List[str]) -> List[List[str]]:
        """Parse markdown table text into a list of rows"""
        table_data = []
        for line in lines:
            # Skip separator lines (containing only |, -, and spaces)
            if re.match(r'^[\|\- ]+$', line):
                continue
            # Split the line into cells and clean them
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty cells from start/end caused by leading/trailing |
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:  # Only add non-empty rows
                table_data.append(cells)
        return table_data

    def _is_table_start(self, line: str) -> bool:
        """Check if a line is the start of a markdown table"""
        return line.strip().startswith('|') or line.strip().startswith('| ')
        
    async def generate_validation_report(
        self,
        validation_id: str,
        tests: List[Dict[str, Any]]
    ) -> str:
        """
        Generate a validation report matching the web client's Test Results tab,
        with proper text and table formatting.
        """
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Model Validation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary section
            doc.add_heading('Executive Summary', level=1)
            summary = doc.add_paragraph()
            summary.add_run('Validation ID: ').bold = True
            summary.add_run(validation_id)
            summary.add_run('\nDate: ').bold = True
            summary.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            
            # Add test results
            doc.add_heading('Test Results', level=1)
            
            for test in tests:
                doc.add_heading(f"Test: {test.get('description', 'Unnamed Test')}", level=2)
                
                if test.get('results'):
                    for result in test['results']:
                        content = result.get('content', '')
                        if not content:
                            continue
                            
                        lines = content.split('\n')
                        i = 0
                        while i < len(lines):
                            line = lines[i].strip()
                            
                            if not line:
                                i += 1
                                continue
                                
                            if line.startswith('!['):
                                # Handle image
                                img_name = line.split('/')[-1].strip(')')
                                img_path = (self.base_path / validation_id / 'tests' / 
                                          f"test_{test.get('test_id')}" / 'images' / img_name)
                                if img_path.exists():
                                    doc.add_picture(str(img_path), width=Inches(6))
                                i += 1
                                
                            elif self._is_table_start(line):
                                # Collect table lines
                                table_lines = []
                                while i < len(lines) and lines[i].strip():
                                    table_lines.append(lines[i])
                                    i += 1
                                
                                # Parse and create table
                                table_data = self._parse_markdown_table(table_lines)
                                if table_data:
                                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                    table.style = 'Table Grid'
                                    # Fill the table
                                    for row_idx, row in enumerate(table_data):
                                        for col_idx, cell_text in enumerate(row):
                                            cell = table.cell(row_idx, col_idx)
                                            cell.text = cell_text.strip()
                                    
                                    # Add spacing after table
                                    doc.add_paragraph()
                                
                            else:
                                # Regular text - remove markdown heading markers
                                clean_text = re.sub(r'^#+\s*', '', line)
                                if clean_text:
                                    doc.add_paragraph(clean_text)
                                i += 1
                
                # Add separator between tests
                doc.add_paragraph('_' * 50)
            
            # Save document
            output_dir = self.base_path / validation_id / 'documentation'
            output_dir.mkdir(exist_ok=True)
            output_path = output_dir / f'validation_report_{validation_id}.docx'
            doc.save(str(output_path))
            
            return str(output_path)
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate report: {str(e)}")